import openai
import streamlit as st
from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
from langchain.document_loaders import PyPDFLoader
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from markdown_pdf import MarkdownPdf, Section
from io import BytesIO
import re
from docx import Document


hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
    """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Configuration de l'API OpenAI

api_key = st.secrets["API_KEY"]
openai.api_key = api_key

def load_and_split_documents(file_path):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
    raw_documents = PyPDFLoader(file_path).load()
    return text_splitter.split_documents(raw_documents)

def create_faiss_db(documents):
    if not documents:
        raise ValueError("Aucun document trouvé pour créer la base de données FAISS.")
    embeddings = OpenAIEmbeddings(openai_api_key=api_key)
    return FAISS.from_documents(documents, embeddings)

def generate_section(system_message, query, documents, combined_content):
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    llm = ChatOpenAI(openai_api_key=api_key)
    if documents:
        db = create_faiss_db(documents)
        qa_chain = ConversationalRetrievalChain.from_llm(llm, retriever=db.as_retriever(), memory=memory, verbose=True)
        combined_info = qa_chain.run({'question': query})
        full_content = combined_content + " " + combined_info + " " + query
    else:
        full_content = combined_content + " " + query
    completion = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": system_message},
            {"role": "user", "content": full_content}
        ],
        temperature=0.9
    )
    return completion['choices'][0]['message']['content']

def extract_company_name(text):
    match = re.search(r"(nom de l'entreprise est|Nom de l'entreprise|La vision de) ([\w\s]+)", text, re.IGNORECASE)
    if match:
        return match.group(2).strip()
    return "Nom de l'entreprise non trouvé"

def generate_markdown(results, company_name):
    markdown_content = "# Business Plan\n\n"
    markdown_content += f"## Entreprise: {company_name}\n\n"
    
    for sec_name, content in results.items():
        markdown_content += f"## {sec_name}\n\n"
        paragraphs = content.split('\n')
        for paragraph in paragraphs:
            if paragraph.startswith('* '):  # Bullet points
                markdown_content += f"- {paragraph[2:]}\n"
            elif paragraph.isdigit() and paragraph[1] == '.' and paragraph[2] == ' ' and paragraph[3] == '*':  # Numbered points with '*'
                markdown_content += f"- {paragraph[4:]}\n"
            else:
                markdown_content += f"{paragraph}\n"
        markdown_content += "\n"

    return markdown_content

def markdown_to_word_via_text(markdown_content):
    # Créer un nouveau document Word
    doc = Document()
    doc.add_heading('Business Plan', 0)

    # Diviser le contenu en lignes
    lines = markdown_content.split('\n')
    table_data = []
    inside_table = False

    for line in lines:
        line = line.strip()
        if not line:
            # Si ligne vide et données de table en cours, ajouter le tableau au document
            if table_data:
                num_cols = len(table_data[0])
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                for i, row in enumerate(table_data):
                    for j, cell in enumerate(row):
                        table.cell(i, j).text = cell.strip()
                table_data = []
                inside_table = False
            continue

        if line.startswith('## '):
            # Sous-titre
            doc.add_heading(line[3:], level=2)
        elif line.startswith('- '):
            # Liste à puces
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            # Liste numérotée
            doc.add_paragraph(line, style='List Number')
        elif line.startswith('|'):
            # Détection des lignes de tableau (évite les lignes de séparation)
            if re.match(r'\|?\s*[-:]+\s*\|', line):
                inside_table = True
                continue  # Ignore les lignes de séparation
            else:
                inside_table = True
                table_data.append([cell.strip() for cell in line.split('|')[1:-1]])  # Enlever les bords vides et espaces
        elif line.startswith('**') and line.endswith('**'):
            # Texte en gras
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[2:-2])
            run.bold = True
        elif not inside_table:
            # Paragraphe normal
            doc.add_paragraph(line)

    # Traiter les données de table restantes
    if table_data:
        num_cols = len(table_data[0])
        table = doc.add_table(rows=len(table_data), cols=num_cols)
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.cell(i, j).text = cell.strip()

    # Sauvegarder le document dans un buffer mémoire
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.title("Ish-AI : Générateur de Business Plan")

    uploaded_file = st.file_uploader("Téléchargez votre fichier PDF", type="pdf")
    user_text_input = st.text_area("Entrez des informations supplémentaires ou un texte alternatif:", height=200)
    
    if uploaded_file or user_text_input:
        documents = []
        combined_content = user_text_input  

    
        if uploaded_file:
            file_path = "uploaded_document.pdf"
            with open(file_path, "wb") as f:
                f.write(uploaded_file.read())
            documents = load_and_split_documents(file_path)    
            # Créer un dictionnaire pour stocker les résultats
        results = {}
            
            # Messages système et requêtes pour chaque section
        system_messages = {
                "Résumé Exécutif": """
                    Générer cette section du business plan:
                    Résumé executif:
                    Generer deux grands paragraphes avec plusieurs lignes:
                    Le résumé exécutif s’écrit tout à la fin, lorsque le business plan est complet. Il s’agit d’une
                    description du projet dans son ensemble, des enseignements du business plan ainsi que du
                    plan financier. Il doit donner l’eau à la bouche au lecteur, qui doit avoir une envie pressante
                    de lire le business plan à la suite du résumé exécutif.
                    Généralement on y injecte un résumé succinct de chaque chapitre et on termine par le plan
                    financier en établissant le montant des financements recherchés afin de lancer le projet.

                    Résumé executif:

                """,
                "Présentation du Projet": """
                    Générer cette section du business plan:
                    1 - PRÉSENTATION DU PROJET

                    1.1. DESCRIPTION DU PRODUIT/SERVICES
                    Générer deux grands paragraphes avec plusieurs lignes :
                    Décrire ici le ou les produit(s) ou service(s) que vous allez proposer à vos consommateurs.
                    Expliquez pourquoi ce produit/service va plaire au public.

                    1.2. ORIGINE DU PROJET
                    Générer deux grands paragraphes avec plusieurs lignes :
                    Décrivez l’origine du projet en quelques phrases. Décrivez le problème que vous avez identifié
                    et comment vous comptez apporter la solution à ce problème. Vous avez donc détecté une
                    opportunité dans le marché. Décrivez comment vous allez la saisir.
                    Il faut convaincre le lecteur que votre solution va répondre à une réelle demande sur le marché
                    et donc que votre projet va réussir. Comment allez-vous convaincre un investisseur qu’il va
                    récupérer le financement qu’il vous aura octroyé ?

                    1.3. LA MISSION, VISION & LES OBJECTIFS DU PROJET
                    Générer deux grands paragraphes avec plusieurs lignes :
                    Votre mission représente la raison d’être de votre entreprise. Comment souhaitez-vous que
                    votre entreprise contribue à la société et à l’économie, mais surtout, qu’entendez-vous
                    proposer à vos clients ?
                    Exemple : « Mon entreprise entend résoudre les problèmes de […] au Nord-Kivu, dans
                    le respect permanent de l’être humain et de l’environnement. »
                    Avant de pouvoir définir vos objectifs, il vous faut d’abord énoncer votre vision. Comment
                    envisagez-vous votre entreprise dans 3 à 5 ans ? Une vision représente donc la direction
                    unique à suivre par votre entreprise.
                    Exemple : « Nous envisageons de développer l’entreprise de manière à ce qu’elle
                    devienne, dans les 5 ans, leader du marché du […] en RDC. »
                    Les objectifs sont des défis que vous allez imposer à votre entreprise afin qu’elle accomplisse
                    sa mission et se dirige vers sa vision. Il peut s’agir d’objectifs qualitatifs (e.g. être rentable au
                    terme de la troisième année, avoir finalisé l’ensemble des installations et aménagements après
                    6 mois, etc.) ou quantitatifs (e.g. vendre 100,000 unités de X au terme de la première
                    année ; avoir employé 5 personnes en 2 mois, etc.). La méthode SMART vous aidera à définir
                    de tels objectifs tangibles.

                    1.4. STADE D’AVANCEMENT DU PROJET & DÉVELOPPEMENTS
                    Générer deux paragraphes de plusieurs lignes :
                    Décrire ce qui a déjà été fait jusqu’à ce jour, les montants qui ont déjà été investis, les
                    prototypes développés ou les tests conduits, le nombre de personnes qui travaillent
                    actuellement sur le projet, vos volumes de ventes actuels, les contrats éventuellement déjà
                    signés, etc.
                    Décrivez également ce qui reste à faire afin de mener le projet à bien. Pour ce faire,
                    établissez un planning de lancement et un calendrier de réalisation : Énumérez les principales
                    étapes de réalisation de votre projet (incorporation ou enregistrement, location d'espaces,
                    achat d'équipements, publicité, date prévue de démarrage, etc.) et définissez à quelle
                    échéance elles auront lieu.

                    Le tableau ci-dessous est un format d'exemple de planning de calendrier de lancement et de réalisation :
                    Générer un diagramme de Gantt qui reprend les activités et associe chaque étape au mois de son exécution, générer aussi les activités et les associer aux mois.
                        M1 M2 M3 M4 M5 M6 M7 M8 M9 M10 M11 M12
                    Étape 1
                    Étape 2
                    Étape 3
                    Étape 4
                    Étape 5
                    Étape 6

                    Générer deux paragraphes de plusieurs lignes pour chaque point.

                    Présentation du projet:
                """,
                "Présentation des Porteurs de Projet": """
                    Générer cette section du business plan:

                    2 - PRÉSENTATION DES PORTEURS DE PROJET

                    2.1 L’ÉQUIPE DE PROJET
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Faites une présentation des fondateurs du projet en soulignant leurs contributions au projet.
                    Mettez en avant les atouts et complémentarités de chacun des promoteurs. Expliquez
                    comment vous vous êtes rencontrés et pourquoi vous avez décidé de vous associer
                    pour ce projet.
                    Mentionnez vos études et formations, vos expériences, savoir-faire technique, organisation du
                    travail, administration d’entreprise, expérience entrepreneuriale, etc. Il ne s’agit pas de présenter
                    votre CV ici ; privilégiez plutôt les informations en relation avec le projet.
                    Rédigez un paragraphe indiquant votre profil, situation familiale, région/ville d’origine. Vous
                    pouvez également mentionner vos motivations pour la création d’entreprise.
                    Indiquez si vous avez des contacts utiles pour le projet dans votre réseau familial et amical, ainsi que les
                    organismes de soutien ou le type d’aide et d’appui qui vous sont disponibles.

                    2.2 CHOIX DE LA FORME JURIDIQUE
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Expliquez sous quelle forme sera enregistrée votre entreprise (Entreprise individuelle, SARL,
                    SA, Société Anonyme, Coopérative...) et en quoi la forme choisie est avantageuse pour
                    vous.
                    Indiquez, en vous basant sur les détails de l’inscription contenus dans les documents, si votre entreprise possède un Numéro d’inscription ou RCCM, ou une patente. Donnez les informations du RCCM ou le Numéro d’inscription que l'entreprise possède. Si nécessaire, joignez une copie
                    de ces documents en annexe du présent document.

                    Présentation des porteurs du projet :

                """,
                "Analyse de Marché": """
                    Générer cette section du business plan:

                    3 – ANALYSE DE MARCHÉ

                    3.1 CARACTÉRISTIQUES DE L’ENVIRONNEMENT/SECTEUR
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Identifiez le secteur d’activité dans lequel œuvrera votre entreprise et décrivez les perspectives
                    d’avenir de ce secteur, ainsi que son environnement. Ne décrivez pas votre entreprise même,
                    mais plutôt l’environnement EXTERNE à votre entreprise (ce qui est commun pour toutes les
                    entreprises œuvrant dans votre secteur). C’est sur la base de ces facteurs externes que vous
                    définirez les « opportunités » et les « menaces » auxquelles fera face l’entreprise que vous
                    souhaitez créer.
                    Le but est de convaincre les lecteurs que vous avez une bonne connaissance de votre secteur
                    et que celui-ci présente assez d’opportunités pour votre projet. Lorsque vous détaillez vos
                    analyses, assurez-vous que l’information que vous présentez soit basée sur des faits, des
                    statistiques, des études et des opinions d’experts. Il convient donc de mentionner vos sources
                    dans la description de votre secteur afin d’être crédible.
                    En plus d’une description générale du contexte dans lequel va opérer votre entreprise, il est
                    important de produire une analyse macroéconomique basée sur le modèle ‘PESTEL’
                    (Politique, Économique, Social, Technologique, Écologique, Légal) et une analyse
                    concurrentielle basée sur les ‘Cinq Forces de Porter’.

                    3.2 MARCHÉ POTENTIEL & MARCHÉ CIBLE
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Il est important de distinguer le marché potentiel (individus ou entreprises susceptibles
                    d’acheter votre produit/service) du marché cible (individus ou entreprises que vous
                    souhaitez viser en particulier dans l’offre de vos produits/services).
                    Le marché potentiel est donc l’ensemble des personnes et des entreprises qui demandent
                    ou qui sont susceptibles de demander vos produits et/ou services pour satisfaire leurs besoins.
                    Ce sont vos consommateurs potentiels. Il doit être divisé en segments de consommateurs
                    (e.g. entreprises/ménages ; ménages à haut/moyen/bas revenus, grandes/moyennes/petites
                    entreprises, personnes âgées/adultes/étudiants/enfants, etc.). Ces segments se caractérisent
                    par leurs préférences et habitudes d’achat (e.g. couleurs, goûts, qualité, budget, accessibilité,
                    etc.). Le marché potentiel se quantifie en nombre d’individus et en volumes potentiellement
                    consommés par ces individus.
                    Sur la base du marché potentiel et des caractéristiques des segments, vous allez définir votre
                    marché cible. Il s’agit donc des segments que vous allez viser en priorité car ils présentent
                    des volumes, des préférences et des capacités de consommation qui correspondent au produit
                    ou service que vous souhaitez offrir.
                    Présentez ici les résultats de votre analyse et tirez-en des conclusions quant au potentiel offert
                    par ce marché. Ensuite, définissez les segments que vous souhaitez cibler et expliquez pourquoi
                    vous avez fait ce choix particulier.

                    3.3 CARACTÉRISTIQUES DE L’OFFRE & CONCURRENCE
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Quels sont les concurrents directs (produits similaires) et indirects (produits substituts) ? Listez
                    les concurrents et décrivez-les. Faites cela pour chacun des produits/services proposés.
                    Listez les points forts, les points faibles, ainsi que les prix pratiqués par la concurrence.
                    Il faudrait générer ce tableau :

                    | Nom          | Forces                       | Faiblesses                   | Prix  |
                    |--------------|------------------------------|------------------------------|-------|
                    | Concurrent 1 |                              |                              |       |
                    | Concurrent 2 |                              |                              |       |
                    | Substitut 1  |                              |                              |       |
                    | Substitut 2  |                              |                              |       |

                    3.4 ANALYSE S.W.O.T.
                    Générer trois grands paragraphes avec plusieurs lignes pour chaque point de l'analyse
                    SWOT (De l'anglais – Strengths : forces, Weaknesses : faiblesses, Opportunities : opportunités,
                    Threats : menaces) est un outil préparatoire à la prise de décision. Il a la particularité d’intégrer
                    les forces et faiblesses propres à l’entreprise ainsi que les opportunités et menaces présentes
                    dans l’environnement dans lequel l’entreprise se trouvera.
                    Veuillez ici présenter les résultats de l’analyse SWOT de votre entreprise.
                    Les forces et les faiblesses décrivent des éléments internes à votre entreprise.
                    - Une force serait par exemple : disposer d’une main-d’œuvre spécialisée et bien formée.
                    - Une faiblesse serait par exemple : un emplacement défavorable car situé à l’écart des
                    principales artères commerciales de la ville.
                    Les opportunités et les menaces décrivent des éléments qui se rapportent à l’environnement
                    externe, c’est-à-dire l’environnement dans lequel vous opérez (celui que vous avez décrit dans la
                    section « Caractéristiques de l’environnement/secteur »).
                    - Une opportunité serait par exemple : Le marché pour votre produit est en plein essor et
                    en pleine croissance.
                    - Une menace serait par exemple : Risque que la concurrence des pays voisins
                    augmente et que des importations à prix réduits viennent inonder le marché.

                """,
                "Moyens de Production et Organisation": """
                    Générer cette section du business plan :

                    4 – MOYENS DE PRODUCTION ET ORGANISATION

                    4.1 PROCESSUS DE PRODUCTION
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Décrivez les étapes de production de vos biens et services de manière détaillée afin que le lecteur comprenne bien ce domaine spécifique. Chaque étape doit être décrite en termes de valeur ajoutée au produit/service.
                    Pour chaque étape, indiquez les intrants (volume de ressources matérielles, temporelles, humaines, etc.) et les extrants (volume des produits finis, déchets, produits dérivés).
                    Une description distincte du processus doit être fournie pour chaque bien et/ou service produit. Vous pouvez utiliser le tableau ci-dessous pour résumer et schématiser le processus de production :

                    | Étape | Intrants | Extrants | RH | Machines & Équipements |
                    |-------|----------|----------|----|------------------------|
                    |       |          |          |    |                        |
                    |       |          |          |    |                        |

                    4.2 CAPACITÉ DE PRODUCTION
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Définissez les limites de production en fonction des ressources matérielles, humaines, temporelles et financières disponibles. Indiquez les facteurs limitants (goulots d’étranglement) de cette production.
                    Complétez le tableau ci-dessous pour détailler les capacités de production :

                    | Étapes | Capacité de Production Maximale par Mois | Facteurs Contraignants |
                    |--------|------------------------------------------|------------------------|
                    |        |                                          |                        |
                    |        |                                          |                        |

                    4.3 HORAIRE DE PRODUCTION
                    Générer le Diagramme de Gantt :
                    Présentez l’horaire de production sur une période donnée avec les étapes et heures correspondantes :

                    Jour 1
                    | Étape | H1 | H2 | H3 | H4 | H5 | H6 | H7 | H8 |
                    |-------|----|----|----|----|----|----|----|----|
                    |       |    |    |    |    |    |    |    |    |

                    Jour 2
                    | Étape | H1 | H2 | H3 | H4 | H5 | H6 | H7 | H8 |
                    |-------|----|----|----|----|----|----|----|----|
                    |       |    |    |    |    |    |    |    |    |

                    Jour 3
                    | Étape | H1 | H2 | H3 | H4 | H5 | H6 | H7 | H8 |
                    |-------|----|----|----|----|----|----|----|----|
                    |       |    |    |    |    |    |    |    |    |

                    Jour 4
                    | Étape | H1 | H2 | H3 | H4 | H5 | H6 | H7 | H8 |
                    |-------|----|----|----|----|----|----|----|----|
                    |       |    |    |    |    |    |    |    |    |

                    Jour 5
                    | Étape | H1 | H2 | H3 | H4 | H5 | H6 | H7 | H8 |
                    |-------|----|----|----|----|----|----|----|----|
                    |       |    |    |    |    |    |    |    |    |

                    4.4 APPROVISIONNEMENT
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Pour chaque matière première et équipement mentionné dans les étapes de production, indiquez le prix d’achat, le fournisseur choisi, les raisons de ce choix, la politique de paiement (cash/crédit, échéance de crédit, etc.), la politique de livraison (enlèvement sur place, livraison gratuite, payante, etc.), les délais de livraison, ainsi que toute autre information pertinente.

                    4.5 LES MOYENS HUMAINS
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Mentionnez le nombre d’emplois créés, excluant ceux des promoteurs. Expliquez les besoins en main-d’œuvre et la rémunération des employés. Précisez si vous prévoyez de coopérer avec des sous-traitants.
                    Présentez, si nécessaire, le diagramme de l’entreprise. Décrivez la structure de votre entreprise en fonction des niveaux de responsabilité des dirigeants et du personnel. Utilisez le tableau ci-dessous pour détailler les fonctions :

                    | Fonction | Nombre | Tâches | Salaire Mensuel |
                    |----------|--------|--------|-----------------|
                    |          |        |        |                 |
                    |          |        |        |                 |

                    4.6 LOCAL & IMPLANTATION
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Définissez les endroits où vous allez implanter vos locaux et installations. Indiquez les raisons stratégiques de votre choix de localisation et énumérez les coûts éventuels de location.

                    4.7 AMÉNAGEMENTS
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Décrivez les aménagements et constructions prévus sur les parcelles que vous allez exploiter. Détaillez les travaux de construction pour aménager votre espace de travail et les coûts associés. Joignez éventuellement un plan d’architecture ou d’aménagement, ainsi que la liste des matériaux que vous utiliserez.
                    Listez les machines et équipements que vous installerez dans vos installations, en précisant leurs caractéristiques techniques (dimensions, puissance électrique, exigences particulières, etc.).

                    4.8 ESTIMATIONS DE VOS COÛTS DE PRODUCTION
                    Générer trois grands paragraphes avec plusieurs lignes :
                    Sur la base des informations fournies ci-dessus, estimez vos coûts de production. Cette estimation doit inclure les frais opérationnels (intrants, matières premières, matériel, salaires, électricité, etc.) nécessaires à la production de vos biens et services, séparément pour chaque bien/service.
                    Tenez compte des coûts variables (liés directement à la quantité produite) et des coûts fixes (loyer, salaires, électricité, etc.) pour obtenir une estimation du coût total de production. Les coûts variables détermineront votre marge brute, qui doit être suffisante pour couvrir vos coûts fixes.
                    Les dépenses d’investissement (terrains, maisons, équipements, machines, etc.) ne doivent pas être incluses dans l’estimation des coûts de production, car elles seront reflétées dans le plan financier sous forme d’amortissements. L’analyse des besoins d’investissement doit être faite séparément.
                    Les résultats de l’analyse des coûts doivent également être intégrés dans le plan financier prévisionnel. Utilisez les tableaux ci-dessous pour compléter l’analyse des coûts variables et fixes :

                    Coûts Variables
                    Listez ici les coûts variables de votre activité, avec une explication et un détail du calcul.

                    | Poste de Coût                          | Montants Mensuels et Annuels (US $) | Explication |
                    |---------------------------------------|-------------------------------------|-------------|
                    | Matières premières (total annuel)     |                                     |             |
                    | Marchandises (total annuel)           |                                     |             |
                    | Autres intrants (total annuel)        |                                     |             |
                    | Matériel (total annuel)               |                                     |             |
                    | Commissions pour agents commerciaux    |                                     |             |
                    | … à compléter                          |                                     |             |

                    Coûts Fixes
                    Listez ici les coûts fixes de votre activité, avec une explication et un détail du calcul.

                    | Poste de Coût                          | Montant Annuel (US $) | Explication |
                    |---------------------------------------|----------------------|-------------|
                    | Assurances                            |                      |             |
                    | Téléphone, internet                   |                      |             |
                    | Autres abonnements                    |                      |             |
                    | Carburant, transports                 |                      |             |
                    | Frais de déplacement et hébergement   |                      |             |
                    | Eau, gaz                              |                      |             |
                    | Fournitures diverses                  |                      |             |
                    | Entretien matériel et vêtements       |                      |             |
                    | Nettoyage des locaux                  |                      |             |
                    | Budget publicité et communication     |                      |             |
                    | Loyer et charges locatives            |                      |             |
                    | Expert-comptable, avocats             |                      |             |
                    | Frais bancaires et terminal carte bleue |                    |             |
                    | Taxes fixes                           |                      |             |
                    | Assurances                            |                      |             |
                    | Électricité Virunga                    |                      |             |
                    | … à compléter                          |                      |             |

                """,
                "Stratégie Marketing et Moyens Commerciaux": """
                Générer cette section du business plan :

                5 – STRATÉGIE MARKETING ET MOYENS COMMERCIAUX

                5.1 MARKETING-MIX
                Le Marketing-Mix est l’ensemble des techniques et des outils utilisés pour assurer que votre produit ou votre service sera attrayant pour vos clients potentiels. Il consiste à définir les caractéristiques de votre offre en fonction de votre marché cible, afin que vos clients se tournent vers vous plutôt que vers la concurrence et afin d’assurer des quantités de ventes suffisantes pour rendre votre projet rentable. Pour attirer vos clients, les inciter à acheter et à revenir régulièrement, vous devez planifier une stratégie efficace qui tient compte des spécificités des segments visés, de leurs préférences et de leur capacité d’achat. Le Marketing-Mix est l’équilibre entre la définition du produit/service, le prix, les moyens de promotion et les moyens de distribution qui convaincront vos clients potentiels.

                1- Politique de produit/service
                Comment allez-vous positionner votre produit/service par rapport à la concurrence afin de vous différencier ?
                Qu’allez-vous proposer au client pour qu’il se tourne vers votre offre plutôt que celle de vos concurrents ou des produits substituts ?
                Qu’est-ce qui fait que votre produit/service est différent des autres et apprécié par les consommateurs ?

                2- Politique de prix
                Quel est le coût de revient de votre produit/service (coûts variables par unité produite, c'est-à-dire le coût additionnel pour produire une unité supplémentaire) ?
                Comment souhaitez-vous vous positionner par rapport au prix pratiqué par la concurrence, au regard de la politique de produit/service décrite ci-dessus ?
                Votre clientèle cible est-elle capable et prête à payer ce prix ?

                3- Politique de distribution
                Comment comptez-vous distribuer vos produits ?
                Où comptez-vous distribuer vos produits ?
                Quels canaux comptez-vous employer ?
                En quoi ce modèle contribuera-t-il à votre réussite ?
                Est-ce différent des méthodes existantes ?
                Cela constitue-t-il un avantage concurrentiel par rapport à vos concurrents ou les produits substituts ?

                4- Politique de communication
                Comment allez-vous véhiculer l’image que vous souhaitez donner à votre produit ?
                - Quel nom, logo et couleurs ?
                - Quel message, slogan ?
                - Quelles actions commerciales et de communication sont prévues dans le temps ?

                Type d’action

                | Type d’action            | Janvier | Février | Mars | … |
                |--------------------------|---------|---------|------|---|
                | Actions pour se faire connaître : |         |         |      |   |
                | - - - -                  | €       | €       | €    | € |
                | Actions pour faire tester ou essayer : |         |         |      |   |
                | - - - -                  | €       | €       | €    | € |
                | Actions pour faire acheter : |         |         |      |   |
                | - - - -                  | €       | €       | €    | € |
                | Actions pour fidéliser : |         |         |      |   |
                | - - - -                  | €       | €       | €    | € |

                Résumez ci-dessous les éléments du marketing-mix par segment visé :

                | Segment de clientèle | Produit/service proposé | Positionnement en terme de prix | Lieu de distribution | Style et mode de communication |
                |----------------------|--------------------------|--------------------------------|----------------------|-------------------------------|
                | Segment 1            |                          |                                |                      |                               |
                | Segment 2            |                          |                                |                      |                               |
                | Segment 3            |                          |                                |                      |                               |

                5.2 PRÉVISIONS DES VENTES
                Entraînez plus en détail. En vous basant sur l’analyse du marché potentiel et votre sélection du segment à cibler, estimez les quantités que vous pourriez vendre. Estimez ensuite les quantités de personnes que vous parviendrez à toucher au sein de votre cible grâce à votre politique de communication et votre modèle de distribution. Enfin, définissez la part des personnes touchées que vous parviendrez à convertir en clients à travers votre politique de produit et de prix. Faites ces estimations pour les quatre années à venir.

                Il est impératif de justifier ces estimations sur la base de suppositions réalistes et/ou de faits déjà acquis (niveau de ventes actuel, contrats de vente signés, etc.).

                L’estimation des quantités prévisionnelles doit être faite pour chaque produit et service séparément.

                Les quantités de ventes que vous estimez doivent tenir compte, entre autres, de votre capacité de production, mais surtout de votre plan de commercialisation, de votre marketing-mix, de votre part de marché espérée, des prix de vente pratiqués par la concurrence et des prix que vous comptez appliquer. Vos prévisions de ventes n’équivaudront donc pas nécessairement à votre capacité de production, mais bien à votre capacité de convaincre vos clients potentiels à acheter votre produit.

                Stratégie Marketing et moyens commerciaux du projet :
                """,
                "Besoin de Démarrage":  """
                    Générer cette section du business plan :

                    6 – BESOIN DE DÉMARRAGE (OU D’INVESTISSEMENT)

                    Veuillez lister ici vos besoins de démarrage (dépenses uniques et généralement amortissables). Il s’agit de vos dépenses d’investissement (comme des terrains, des maisons, des équipements, des machines, etc.). Notez que ces dépenses peuvent déjà avoir été encourues et qu’elles ne sont donc plus à prévoir.

                    Veuillez aussi expliquer comment vous comptez financer ces besoins d’investissements :
                    - Apport sur fonds propres
                    - Apport personnel
                    - Prêt à la banque
                    - Prêt Virunga
                    - Autres sources

                    Expliquez la justification et le détail du calcul. Commentez les chiffres.

                    | Poste de coût                            | Montant (US $) | Explication                           |
                    |------------------------------------------|----------------|---------------------------------------|
                    | Frais d’établissement                    |                |                                       |
                    | Frais d’ouverture de compteurs           |                |                                       |
                    | Logiciels, formations                    |                |                                       |
                    | Dépôt marque, brevet, modèle             |                |                                       |
                    | Droits d’entrée                          |                |                                       |
                    | Achat fonds de commerce ou parts          |                |                                       |
                    | Droit au bail                             |                |                                       |
                    | Caution ou dépôt de garantie              |                |                                       |
                    | Frais de notaire ou d’avocat              |                |                                       |
                    | Enseigne et éléments de communication    |                |                                       |
                    | Achat immobilier                         |                |                                       |
                    | Travaux et aménagements                  |                |                                       |
                    | Matériel/Machines/Équipements            |                |                                       |
                    | Matériel de bureau                       |                |                                       |
                    | Stock de matières et produits pour démarrage |              |                                       |
                    | Trésorerie de départ                     |                |                                       |
                    | … à compléter                            |                |                                       |

                    Besoin de démarrage du projet :
                """,
                "Annexes": """
                Générer cette section du business plan:

                7 – ANNEXES
                Renvoyer en annexe les documents trop volumineux ou difficiles à lire : - - - -
                étude de marché complète,
                contrats,
                conditions

                Annexes du projet:

                """
            }

        queries = {
                "Résumé Exécutif": "Générer un résumé exécutif pour cette entreprise.",
                "Présentation du Projet": "Présenter le projet en détail.",
                "Présentation des Porteurs de Projet": "Décrire les membres de l'équipe et leurs qualifications.",
                "Analyse de Marché": "Analyser le marché cible pour cette entreprise.",
                "Moyens de Production et Organisation": "Décrire les moyens de production et l'organisation opérationnelle de cette entreprise.",
                "Stratégie Marketing et Moyens Commerciaux": "Élaborer un plan marketing détaillé pour cette entreprise.",
                "Besoin de Démarrage": "Décrire les besoins en investissement pour démarrer cette entreprise.",
                "Annexes": "Inclure les documents annexes pertinents pour cette entreprise."
            }

        # Espaces réservés pour chaque section
        placeholders = {name: st.empty() for name in system_messages.keys()}

        # Générer toutes les sections automatiquement
        for section_name in system_messages.keys():
            with st.spinner(f"Génération de {section_name}..."):
                system_message = system_messages[section_name]
                query = queries[section_name]
                try:
                    results[section_name] = generate_section(system_message, query, documents, combined_content)
                except ValueError as e:
                    results[section_name] = f"Erreur: {str(e)}"
                combined_content += " " + results[section_name]
                placeholders[section_name].markdown(f"**{section_name}**\n\n{results[section_name]}")
            
        # Extraction du nom de l'entreprise
        first_section_content = results.get("Résumé Exécutif", "")
        company_name = extract_company_name(first_section_content)

        # Créer le contenu Markdown et offrir le téléchargement
        markdown_content = generate_markdown(results, company_name)
        pdf = MarkdownPdf(toc_level=2)
        pdf.add_section(Section(markdown_content))
        pdf.meta["title"] = "Business Plan"
        pdf.meta["author"] = company_name
        pdf_file_path = "business_plan.pdf"
        pdf.save(pdf_file_path)

        # Créer le document Word
        word_buffer =  markdown_to_word_via_text(markdown_content)

        st.success("Le PDF et le document Word ont été générés avec succès.")
        with open(pdf_file_path, "rb") as f:
            st.download_button("Téléchargez le PDF", f, file_name="business_plan.pdf", mime="application/pdf")
            
        st.download_button("Téléchargez le document Word", word_buffer, file_name="business_plan.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Veuillez soumettre un fichier PDF, saisir du texte, ou les deux pour générer un business plan.")


if __name__ == "__main__":
    main()

